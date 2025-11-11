"""
Script to read Word (.docx) and Excel (.xlsx) files
"""
import sys
from pathlib import Path

def read_docx(file_path):
    """Read Word document and return text content"""
    try:
        from docx import Document
        doc = Document(file_path)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Also read tables if any
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_data))
        
        return "\n".join(content)
    except ImportError:
        return "Error: python-docx library not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error reading Word document: {str(e)}"

def read_xlsx(file_path):
    """Read Excel file and return content"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        content = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            content.append(f"\n{'='*50}")
            content.append(f"SHEET: {sheet_name}")
            content.append('='*50 + "\n")
            
            for row in sheet.iter_rows(values_only=True):
                # Skip completely empty rows
                if any(cell is not None and str(cell).strip() for cell in row):
                    row_str = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    content.append(row_str)
        
        return "\n".join(content)
    except ImportError:
        return "Error: openpyxl library not installed. Run: pip install openpyxl"
    except Exception as e:
        return f"Error reading Excel document: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_documents.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == ".docx":
        content = read_docx(file_path)
    elif file_ext in [".xlsx", ".xls"]:
        content = read_xlsx(file_path)
    else:
        content = f"Unsupported file type: {file_ext}"
    
    print(content)

